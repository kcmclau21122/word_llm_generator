"""
Table calculator module for detecting and performing calculations in Word tables.
Identifies Total and Difference labels and calculates appropriate values.
"""

import re
from typing import List, Optional, Tuple
from docx import Document
from docx.table import Table, _Cell

from src.logger_setup import get_logger


class TableCalculator:
    """
    Performs automated calculations on Word document tables.
    Detects calculation requirements from labels and computes results.
    """
    
    def __init__(self):
        """
        Initialize table calculator.
        """
        self.logger = get_logger()
        
        # Patterns for detecting calculation labels
        self.total_pattern = re.compile(r'\b(total|sum)\b', re.IGNORECASE)
        self.difference_pattern = re.compile(r'\b(diff(?:erence)?|delta|variance)\b', re.IGNORECASE)
        self.average_pattern = re.compile(r'\b(avg|average|mean)\b', re.IGNORECASE)
    
    def process_all_tables(self, doc: Document) -> int:
        """
        Process all tables in document and perform calculations.
        
        Args:
            doc: Document object
            
        Returns:
            Number of calculations performed
        """
        if not doc.tables:
            self.logger.info("No tables found in document")
            return 0
        
        self.logger.info(f"Processing {len(doc.tables)} table(s)")
        total_calculations = 0
        
        for table_idx, table in enumerate(doc.tables):
            self.logger.debug(f"Processing table {table_idx + 1}")
            calculations = self._process_table(table)
            total_calculations += calculations
        
        self.logger.info(f"Completed {total_calculations} calculation(s) across all tables")
        return total_calculations
    
    def _process_table(self, table: Table) -> int:
        """
        Process a single table and perform calculations.
        
        Args:
            table: Table object
            
        Returns:
            Number of calculations performed
        """
        calculations = 0
        row_count = len(table.rows)
        col_count = len(table.columns) if table.rows else 0
        
        # Check each cell for calculation labels
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                label = cell.text.strip().lower()
                
                # Check for total
                if self.total_pattern.search(label):
                    if self._perform_calculation(table, row_idx, col_idx, "total"):
                        calculations += 1
                
                # Check for difference
                elif self.difference_pattern.search(label):
                    if self._perform_calculation(table, row_idx, col_idx, "difference"):
                        calculations += 1
                
                # Check for average
                elif self.average_pattern.search(label):
                    if self._perform_calculation(table, row_idx, col_idx, "average"):
                        calculations += 1
        
        return calculations
    
    def _perform_calculation(
        self,
        table: Table,
        label_row: int,
        label_col: int,
        calc_type: str
    ) -> bool:
        """
        Perform calculation based on label position and type.
        
        Args:
            table: Table object
            label_row: Row index of label
            label_col: Column index of label
            calc_type: Type of calculation ("total", "difference", "average")
            
        Returns:
            True if calculation performed successfully
        """
        try:
            # Determine if this is a row or column calculation
            # Labels in first column typically mean row calculation
            # Labels in bottom row typically mean column calculation
            
            is_row_calculation = label_col == 0
            is_col_calculation = label_row == len(table.rows) - 1
            
            if is_row_calculation and label_col + 1 < len(table.columns):
                # Row calculation: sum/calculate across the row
                return self._calculate_row(table, label_row, label_col, calc_type)
            
            elif is_col_calculation and label_row > 0:
                # Column calculation: sum/calculate down the column
                return self._calculate_column(table, label_row, label_col, calc_type)
            
            else:
                self.logger.debug(f"Ambiguous calculation position at ({label_row}, {label_col})")
                return False
                
        except Exception as e:
            self.logger.error(f"Calculation failed at ({label_row}, {label_col}): {str(e)}")
            return False
    
    def _calculate_row(
        self,
        table: Table,
        row_idx: int,
        start_col: int,
        calc_type: str
    ) -> bool:
        """
        Calculate values across a row.
        
        Args:
            table: Table object
            row_idx: Row index
            start_col: Starting column (after label)
            calc_type: Calculation type
            
        Returns:
            True if successful
        """
        row = table.rows[row_idx]
        values = []
        
        # Extract numeric values from cells after the label
        for col_idx in range(start_col + 1, len(row.cells)):
            value = self._extract_number(row.cells[col_idx].text)
            if value is not None:
                values.append(value)
        
        if not values:
            self.logger.debug(f"No numeric values found in row {row_idx}")
            return False
        
        # Perform calculation
        if calc_type == "total":
            result = sum(values)
        elif calc_type == "average":
            result = sum(values) / len(values)
        elif calc_type == "difference":
            result = values[0] - sum(values[1:]) if len(values) > 1 else values[0]
        else:
            return False
        
        # Find result cell (typically last cell in row)
        result_col = len(row.cells) - 1
        self._set_cell_value(row.cells[result_col], result, row.cells[start_col + 1].text)
        
        self.logger.info(f"Row {row_idx} {calc_type}: {result}")
        return True
    
    def _calculate_column(
        self,
        table: Table,
        row_idx: int,
        col_idx: int,
        calc_type: str
    ) -> bool:
        """
        Calculate values down a column.
        
        Args:
            table: Table object
            row_idx: Row index (typically last row)
            col_idx: Column index
            calc_type: Calculation type
            
        Returns:
            True if successful
        """
        values = []
        
        # Extract numeric values from cells above the label
        for r_idx in range(row_idx):
            cell_text = table.rows[r_idx].cells[col_idx].text
            value = self._extract_number(cell_text)
            if value is not None:
                values.append(value)
        
        if not values:
            self.logger.debug(f"No numeric values found in column {col_idx}")
            return False
        
        # Perform calculation
        if calc_type == "total":
            result = sum(values)
        elif calc_type == "average":
            result = sum(values) / len(values)
        elif calc_type == "difference":
            result = values[0] - sum(values[1:]) if len(values) > 1 else values[0]
        else:
            return False
        
        # Set result in the label cell
        result_cell = table.rows[row_idx].cells[col_idx]
        original_format = table.rows[row_idx - 1].cells[col_idx].text if row_idx > 0 else ""
        self._set_cell_value(result_cell, result, original_format)
        
        self.logger.info(f"Column {col_idx} {calc_type}: {result}")
        return True
    
    def _extract_number(self, text: str) -> Optional[float]:
        """
        Extract numeric value from cell text.
        Handles currency symbols, commas, percentages.
        
        Args:
            text: Cell text
            
        Returns:
            Numeric value or None if not found
        """
        if not text:
            return None
        
        # Remove currency symbols and whitespace
        cleaned = re.sub(r'[$€£¥,\s]', '', text)
        
        # Check for percentage
        is_percentage = '%' in text
        cleaned = cleaned.replace('%', '')
        
        # Extract number using regex
        match = re.search(r'-?\d+\.?\d*', cleaned)
        if match:
            value = float(match.group())
            if is_percentage:
                value = value / 100
            return value
        
        return None
    
    def _set_cell_value(self, cell: _Cell, value: float, format_reference: str) -> None:
        """
        Set cell value with appropriate formatting.
        
        Args:
            cell: Cell object
            value: Numeric value
            format_reference: Reference text for format detection
        """
        # Detect format from reference
        has_currency = bool(re.search(r'[$€£¥]', format_reference))
        is_percentage = '%' in format_reference
        
        # Format the value
        if is_percentage:
            formatted = f"{value * 100:.1f}%"
        elif has_currency:
            # Detect currency symbol
            currency_match = re.search(r'([$€£¥])', format_reference)
            symbol = currency_match.group(1) if currency_match else '$'
            formatted = f"{symbol}{value:,.2f}"
        else:
            # Check decimal places in reference
            if '.' in format_reference:
                formatted = f"{value:,.2f}"
            else:
                formatted = f"{value:,.0f}"
        
        # Set cell text
        cell.text = formatted
        self.logger.debug(f"Set cell value: {formatted}")
    
    def validate_table_structure(self, table: Table) -> bool:
        """
        Validate that table has consistent structure for calculations.
        
        Args:
            table: Table object
            
        Returns:
            True if table structure is valid
        """
        if not table.rows:
            return False
        
        # Check that all rows have same number of columns
        col_counts = [len(row.cells) for row in table.rows]
        if len(set(col_counts)) > 1:
            self.logger.warning("Table has inconsistent column counts")
            return False
        
        return True