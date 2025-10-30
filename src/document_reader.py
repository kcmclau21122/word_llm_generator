"""
Document reader module for extracting sections and structure from Word documents.
Identifies headings, section content, and placeholders for content insertion.
"""

from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional
from docx import Document
from docx.text.paragraph import Paragraph

from src.logger_setup import get_logger


@dataclass
class Section:
    """
    Represents a document section with heading and content.
    """
    title: str
    level: int
    content_paragraphs: List[str]
    has_placeholder: bool
    placeholder_index: Optional[int]
    paragraph_index: int


class DocumentReader:
    """
    Reads Word documents and extracts structured section information.
    Identifies headings and content placeholders for LLM generation.
    """
    
    def __init__(self, heading_styles: Optional[List[str]] = None, placeholder_pattern: str = "{{SECTION_CONTENT}}"):
        """
        Initialize document reader.
        
        Args:
            heading_styles: List of paragraph style names to treat as section headings
            placeholder_pattern: Pattern to identify content placeholders
        """
        self.heading_styles = heading_styles or ["Heading 1", "Heading 2", "Heading 3"]
        self.placeholder_pattern = placeholder_pattern
        self.logger = get_logger()
    
    def load_document(self, file_path: str) -> Document:
        """
        Load a Word document from file.
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            Document object
            
        Raises:
            FileNotFoundError: If file doesn't exist
            Exception: If document cannot be loaded
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        self.logger.info(f"Loading document: {file_path}")
        
        try:
            doc = Document(file_path)
            self.logger.info(f"Document loaded successfully: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
            return doc
        except Exception as e:
            self.logger.error(f"Failed to load document: {str(e)}")
            raise
    
    def extract_sections(self, doc: Document) -> List[Section]:
        """
        Extract sections from document based on heading styles.
        
        Args:
            doc: Document object
            
        Returns:
            List of Section objects
        """
        self.logger.info("Extracting sections from document")
        sections: List[Section] = []
        current_section: Optional[Section] = None
        
        for idx, paragraph in enumerate(doc.paragraphs):
            style_name = paragraph.style.name if paragraph.style else ""
            
            # Check if this is a heading
            if self._is_heading(style_name):
                # Save previous section if exists
                if current_section is not None:
                    sections.append(current_section)
                    self.logger.debug(f"Section added: {current_section.title}")
                
                # Start new section
                heading_level = self._get_heading_level(style_name)
                current_section = Section(
                    title=paragraph.text.strip(),
                    level=heading_level,
                    content_paragraphs=[],
                    has_placeholder=False,
                    placeholder_index=None,
                    paragraph_index=idx
                )
                self.logger.debug(f"New section started: {current_section.title} (Level {heading_level})")
            
            elif current_section is not None:
                # Add paragraph to current section
                text = paragraph.text.strip()
                if text:
                    current_section.content_paragraphs.append(text)
                    
                    # Check for placeholder
                    if self.placeholder_pattern in text:
                        current_section.has_placeholder = True
                        current_section.placeholder_index = len(current_section.content_paragraphs) - 1
                        self.logger.debug(f"Placeholder found in section: {current_section.title}")
        
        # Add last section
        if current_section is not None:
            sections.append(current_section)
        
        self.logger.info(f"Extracted {len(sections)} sections")
        return sections
    
    def _is_heading(self, style_name: str) -> bool:
        """
        Check if a style name represents a heading.
        
        Args:
            style_name: Paragraph style name
            
        Returns:
            True if style is a heading
        """
        return style_name in self.heading_styles
    
    def _get_heading_level(self, style_name: str) -> int:
        """
        Extract heading level from style name.
        
        Args:
            style_name: Paragraph style name
            
        Returns:
            Heading level (1, 2, 3, etc.)
        """
        # Try to extract number from style name (e.g., "Heading 1" -> 1)
        for char in style_name.split():
            if char.isdigit():
                return int(char)
        return 1  # Default to level 1
    
    def get_sections_needing_content(self, sections: List[Section]) -> List[Section]:
        """
        Filter sections that have placeholders and need content generation.
        
        Args:
            sections: List of all sections
            
        Returns:
            List of sections with placeholders
        """
        sections_with_placeholders = [s for s in sections if s.has_placeholder]
        self.logger.info(f"Found {len(sections_with_placeholders)} sections needing content")
        return sections_with_placeholders
    
    def get_section_context(self, sections: List[Section], target_section: Section, context_window: int = 2) -> str:
        """
        Get contextual information from surrounding sections for better generation.
        
        Args:
            sections: All document sections
            target_section: Section to get context for
            context_window: Number of previous sections to include
            
        Returns:
            Formatted context string
        """
        try:
            target_idx = sections.index(target_section)
            start_idx = max(0, target_idx - context_window)
            
            context_sections = sections[start_idx:target_idx]
            
            context_parts = []
            for sec in context_sections:
                content_preview = " ".join(sec.content_paragraphs[:2])[:200]
                context_parts.append(f"Previous section '{sec.title}': {content_preview}...")
            
            context = "\n".join(context_parts)
            self.logger.debug(f"Built context for section '{target_section.title}': {len(context)} chars")
            return context
        except ValueError:
            self.logger.warning(f"Section not found in list: {target_section.title}")
            return ""
    
    def get_document_info(self, doc: Document) -> dict:
        """
        Extract general document information.
        
        Args:
            doc: Document object
            
        Returns:
            Dictionary with document metadata
        """
        info = {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "section_count": len(doc.sections)
        }
        
        # Try to get document properties
        try:
            if hasattr(doc.core_properties, "title") and doc.core_properties.title:
                info["title"] = doc.core_properties.title
            if hasattr(doc.core_properties, "author") and doc.core_properties.author:
                info["author"] = doc.core_properties.author
        except Exception:
            pass
        
        self.logger.info(f"Document info: {info}")
        return info