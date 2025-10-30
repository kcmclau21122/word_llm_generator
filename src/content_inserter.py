"""
Content inserter module for adding generated content to Word documents.
Handles placeholder replacement while preserving document formatting.
"""

from pathlib import Path
from typing import List
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from src.document_reader import Section
from src.logger_setup import get_logger


class ContentInserter:
    """
    Inserts generated content into Word documents at placeholder locations.
    Attempts to preserve formatting while adding new content.
    """
    
    def __init__(self, placeholder_pattern: str = "{{SECTION_CONTENT}}"):
        """
        Initialize content inserter.
        
        Args:
            placeholder_pattern: Pattern identifying content placeholders
        """
        self.placeholder_pattern = placeholder_pattern
        self.logger = get_logger()
    
    def insert_content(
        self,
        doc: Document,
        section: Section,
        generated_content: str,
        preserve_placeholder: bool = False
    ) -> bool:
        """
        Insert generated content into document at section placeholder.
        
        Args:
            doc: Document object
            section: Section containing placeholder
            generated_content: Content to insert
            preserve_placeholder: If True, keeps placeholder for reference
            
        Returns:
            True if insertion successful, False otherwise
        """
        if not section.has_placeholder:
            self.logger.warning(f"Section '{section.title}' has no placeholder")
            return False
        
        self.logger.info(f"Inserting content into section: {section.title}")
        
        try:
            # Find the paragraph containing the placeholder
            placeholder_para = None
            para_index = -1
            
            for idx, para in enumerate(doc.paragraphs):
                if self.placeholder_pattern in para.text:
                    placeholder_para = para
                    para_index = idx
                    break
            
            if placeholder_para is None:
                self.logger.error(f"Placeholder not found in document for section: {section.title}")
                return False
            
            # Get the style of the paragraph for formatting consistency
            original_style = placeholder_para.style
            
            # Split content into paragraphs
            content_paragraphs = self._split_into_paragraphs(generated_content)
            
            # Replace placeholder text or insert after
            if preserve_placeholder:
                # Insert content after placeholder
                self._insert_after_paragraph(doc, para_index, content_paragraphs, original_style)
            else:
                # Replace placeholder with first paragraph
                placeholder_para.clear()
                placeholder_para.add_run(content_paragraphs[0])
                placeholder_para.style = original_style
                
                # Insert remaining paragraphs
                if len(content_paragraphs) > 1:
                    self._insert_after_paragraph(doc, para_index, content_paragraphs[1:], original_style)
            
            self.logger.info(f"Successfully inserted {len(content_paragraphs)} paragraph(s)")
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to insert content: {str(e)}")
            return False
    
    def _split_into_paragraphs(self, content: str) -> List[str]:
        """
        Split generated content into separate paragraphs.
        
        Args:
            content: Content text
            
        Returns:
            List of paragraph strings
        """
        # Split on double newlines or single newlines followed by significant content
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
        
        # If no double newlines, split on single newlines
        if len(paragraphs) == 1:
            paragraphs = [p.strip() for p in content.split("\n") if p.strip()]
        
        self.logger.debug(f"Split content into {len(paragraphs)} paragraph(s)")
        return paragraphs
    
    def _insert_after_paragraph(
        self,
        doc: Document,
        after_index: int,
        paragraphs: List[str],
        style
    ) -> None:
        """
        Insert multiple paragraphs after a specific index.
        
        Args:
            doc: Document object
            after_index: Index to insert after
            paragraphs: List of paragraph texts
            style: Paragraph style to apply
        """
        # Get the paragraph element
        base_para = doc.paragraphs[after_index]
        
        for i, para_text in enumerate(paragraphs):
            # Create new paragraph using XML manipulation
            new_para = doc.add_paragraph(para_text, style=style)
            
            # Move the new paragraph to the correct position
            # Get the last paragraph element and move it
            last_para_element = doc.paragraphs[-1]._element
            base_para._element.addnext(last_para_element)
            
            # Update base_para for next iteration
            base_para = doc.paragraphs[after_index + i + 1]
    
    def replace_all_placeholders(
        self,
        doc: Document,
        content_map: dict[str, str]
    ) -> int:
        """
        Replace all placeholders in document with provided content.
        
        Args:
            doc: Document object
            content_map: Dictionary mapping placeholder text to replacement content
            
        Returns:
            Number of replacements made
        """
        replacements = 0
        
        for para in doc.paragraphs:
            for placeholder, content in content_map.items():
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, content)
                    replacements += 1
                    self.logger.debug(f"Replaced placeholder: {placeholder}")
        
        self.logger.info(f"Made {replacements} placeholder replacement(s)")
        return replacements
    
    def save_document(self, doc: Document, output_path: str) -> bool:
        """
        Save document to file.
        
        Args:
            doc: Document object
            output_path: Path to save document
            
        Returns:
            True if save successful, False otherwise
        """
        try:
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            self.logger.info(f"Saving document to: {output_path}")
            doc.save(output_path)
            self.logger.info("Document saved successfully")
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to save document: {str(e)}")
            return False
    
    def create_output_filename(self, original_filename: str, suffix: str = "_generated") -> str:
        """
        Create output filename based on original with suffix.
        
        Args:
            original_filename: Original file name
            suffix: Suffix to add before extension
            
        Returns:
            New filename
        """
        path = Path(original_filename)
        new_name = f"{path.stem}{suffix}{path.suffix}"
        self.logger.debug(f"Generated output filename: {new_name}")
        return new_name